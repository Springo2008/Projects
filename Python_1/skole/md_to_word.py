"""
Konverterer markdown til Word dokument
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Opprett nytt Word dokument
doc = Document()

# Les markdown filen
with open('/Users/husby/Documents/GitHub/Projects/Python_1/skole/ANALYSE_OG_TOLKNING.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Parse markdown og lag Word dokument
lines = content.split('\n')
for line in lines:
    # Sjekk for overskrifter
    if line.startswith('# '):
        # Hovedtittel
        heading = doc.add_heading(line[2:], level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True
    elif line.startswith('## '):
        # Underoverskrift
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        # Suboverskrift
        doc.add_heading(line[4:], level=2)
    elif line.startswith('- '):
        # Punktliste
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('| '):
        # Tabell - finnes ikke her, skip
        pass
    elif line.strip() == '':
        # Tom linje
        doc.add_paragraph()
    else:
        # Normal tekst
        if line.strip():
            doc.add_paragraph(line)

# Håndter tabeller manuelt
# Tabell 1: Deskriptiv statistikk
doc.add_heading('Deskriptiv statistikk', level=2)
table = doc.add_table(rows=4, cols=5)
table.style = 'Light Grid Accent 1'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = 'Lokasjon'
header_cells[1].text = 'Gjennomsnitt'
header_cells[2].text = 'Standardavvik'
header_cells[3].text = 'Min'
header_cells[4].text = 'Max'

# Data rows
data = [
    ['Mitt rom (DOWN)', '24,60°C', '0,59°C', '23°C', '25°C'],
    ['3. etasje stua (UP)', '20,01°C', '0,10°C', '20°C', '21°C'],
    ['Forskjell', '+4,59°C', '-', '-', '-']
]

for i, row_data in enumerate(data):
    row = table.rows[i+1]
    for j, cell_data in enumerate(row_data):
        row.cells[j].text = cell_data

# Legg til resten av innholdet
doc.add_heading('Statistisk test', level=2)
p = doc.add_paragraph()
p.add_run('T-statistikk: ').bold = True
p.add_run('-74,61\n')
p.add_run('P-verdi: ').bold = True
p.add_run('< 0,001 (p < 0,0001)\n')
p.add_run('Konklusjon: ').bold = True
p.add_run('✓✓✓ SVÆRT SIGNIFIKANT FORSKJELL')

doc.add_paragraph()
doc.add_paragraph('Denne sterke p-verdien betyr at det er mindre enn 0,01% sannsynlighet for at denne temperaturforskjellen skyldes tilfeldig variasjon.')

# Lagre dokument
doc.save('/Users/husby/Documents/GitHub/Projects/Python_1/skole/ANALYSE_OG_TOLKNING.docx')
print("✓ Word dokument lagret: ANALYSE_OG_TOLKNING.docx")
